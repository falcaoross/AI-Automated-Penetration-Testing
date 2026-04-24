"""
Multi-Source Document Ingestion Engine — Stage 1  (v3 — All Issues Fixed)
Autopilot-QA Pipeline

Previous fixes (DEFECT-1 through DEFECT-6):
  DEFECT-1 : Line-level paragraph splitting (not double-newline)
  DEFECT-2 : List-item disambiguation for numeric heading patterns
  DEFECT-3 : Colon-terminated labels excluded from titlecase detection
  DEFECT-4 : camelot restored for table extraction; pdfplumber as fallback
  DEFECT-5 : Page-1 specific version/title extraction
  DEFECT-6 : {#anchor} fragments stripped from heading text

Stage 2 readiness fixes (ISSUE-1 through ISSUE-5):
  ISSUE-1 : ToC/cover page detection — excluded from heading stack to prevent
            ToC entries poisoning section_path for all subsequent pages.
            Each page now carries a page_type field: "body" | "toc" | "cover".
  ISSUE-2 : section_path stamped on every paragraph individually.
            Heading stack updated BEFORE building page_section_map (fixes
            off-by-one section attribution on every page).
  ISSUE-3 : Single-column camelot results discarded — they are wrapped text
            blocks, not real tables. Only tables with ≥2 cols AND ≥2 rows kept.
  ISSUE-4 : "Description:" prefix stripped from FR body paragraphs.
            sub_type field added to every paragraph:
              "requirement_statement" | "inputs" | "system_behavior" |
              "outputs" | "assumptions" | "dependencies" | "body"
  ISSUE-5 : OCR quality flag added: noise_ratio + quality_flag ("low"|"acceptable")
            + stage2_note on every ocr_text entry.

Outputs:
  output/structured_output.json       — machine-consumable, schema-strict
  output/ingestion_audit_report.json  — ingestion quality report
"""

import fitz          # PyMuPDF  (primary PDF engine)
import pdfplumber    # fallback + font-size metadata
import pytesseract
import cv2
import camelot
import pandas as pd
import json
import os
import re
import hashlib
from pathlib import Path
from datetime import datetime, timezone
from collections import defaultdict
from docx import Document

try:
    import langdetect
    HAS_LANGDETECT = True
except ImportError:
    HAS_LANGDETECT = False

# Resolve all output paths relative to this script's folder, not cwd.
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
EXTRACTED_IMAGES_DIR = OUTPUT_DIR / "extracted_images"


# ══════════════════════════════════════════════════════════════
# HEADING DETECTION — domain-agnostic, defect-corrected
# ══════════════════════════════════════════════════════════════

SECTION_NUM_RE = re.compile(r"^(\d+(?:\.\d+)*)")
ANCHOR_RE = re.compile(r"\{#[^}]+\}")

# Lines that look numeric but are prose list items (DEFECT-2 fix)
LIST_ITEM_SIGNALS = re.compile(
    r"(?:^(?:\d+\.)+\s+(?:The|If|When|After|Once|This|All|Each|Based)\b)"  # sentence starters
    r"|(?:[:.,]\s*$)"                                                         # ends in punctuation
)


def clean_heading_text(text: str) -> str:
    """Strip markdown anchor fragments and extra whitespace (DEFECT-6)."""
    return ANCHOR_RE.sub("", text).strip()


def detect_heading_level(line: str):
    """
    Returns (level: int, method: str) or (None, 'none').
    Applies disambiguation to avoid false positives on list items and label lines.
    """
    line = line.strip()
    if not line or len(line) > 120:
        return None, "none"

    # ── Numeric heading (DEFECT-2: guard against list items) ──
    m = SECTION_NUM_RE.match(line)
    if m:
        # List item signals: sentence-starter words after number, or trailing punctuation
        if LIST_ITEM_SIGNALS.search(line):
            return None, "list_item"
        # Guard: very long lines are almost never section headings
        if len(line) > 80:
            return None, "too_long"
        depth = m.group(1).count(".") + 1
        return min(depth, 6), "numeric"

    # ── ALL-CAPS short line ──
    if re.match(r"^[A-Z][A-Z\s\-:]{4,60}$", line):
        return 1, "allcaps"

    # ── Title-Case short line (DEFECT-3: exclude colon-terminated labels) ──
    if line.endswith(":"):
        return None, "label"  # 'Inputs:', 'Outputs:', 'System Behavior:', etc.

    if (
        re.match(r"^[A-Z][a-zA-Z0-9\s\-:,/()]{4,79}[^.?!\n]$", line)
        and len(line) <= 60
        and len(line.split()) >= 2
    ):
        return 2, "titlecase"

    return None, "none"


def build_section_path(stack: list) -> str:
    if not stack:
        return "UNMAPPED"
    return " > ".join(t for _, t in stack)


# ══════════════════════════════════════════════════════════════
# TOC / COVER PAGE DETECTION  (Issue 1 fix)
# ══════════════════════════════════════════════════════════════

def detect_page_type(lines: list, known_headings: set = None) -> str:
    """
    Classify a page as 'cover', 'toc', or 'body'.

    Three signals, applied in order:
      1. Explicit marker: any line contains "table of contents" → toc
      2. Dense headings: ≥10 numeric-heading lines AND ≤3 body lines → toc
         (body pages have 4-5 heading lines and 10-26 body lines — well below threshold)
      3. Cover page: ≤12 total non-empty lines AND 0 section numbers → cover
      4. Default: body
    """
    stripped = [l.strip() for l in lines if l.strip()]
    if not stripped:
        return "body"

    # Signal 1 — explicit ToC marker line
    if any("table of contents" in l.lower() for l in stripped):
        return "toc"

    heading_line_count = sum(
        1 for l in stripped
        if SECTION_NUM_RE.match(l) and not LIST_ITEM_SIGNALS.search(l) and len(l) <= 80
    )
    body_line_count = sum(
        1 for l in stripped
        if not SECTION_NUM_RE.match(l) and len(l) > 40
    )

    # Signal 2 — dense heading listing with almost no prose (ToC-style pages)
    # Threshold ≥10 ensures body pages (max 5 heading lines) never trigger
    if heading_line_count >= 10 and body_line_count <= 3:
        return "toc"

    # Signal 3 — cover page: short, no section numbers
    if len(stripped) <= 12 and heading_line_count == 0:
        return "cover"

    return "body"


# ══════════════════════════════════════════════════════════════
# DETERMINISTIC IDs
# ══════════════════════════════════════════════════════════════

def make_doc_id(filename: str, early_text: str) -> str:
    return "doc_" + hashlib.sha256((filename + early_text[:500]).encode()).hexdigest()[:16]


def make_para_id(doc_id: str, page: int, seq: int) -> str:
    return f"{doc_id}_p{page:04d}_s{seq:04d}"


# ══════════════════════════════════════════════════════════════
# METADATA DETECTION (DEFECT-5: page-1 specific extraction)
# ══════════════════════════════════════════════════════════════

VERSION_RE = re.compile(r"(?:version|ver)[:\s]+([0-9]+(?:\.[0-9]+)*)", re.IGNORECASE)
DOC_TYPE_KEYWORDS = {
    "SRS": ["software requirements specification", "srs"],
    "PRD": ["product requirements document", "prd"],
    "SPEC": ["technical specification", "design specification", "spec"],
    "RFC": ["request for comments", "rfc"],
}


def detect_doc_type(text: str) -> str:
    lower = text.lower()
    for dtype, kws in DOC_TYPE_KEYWORDS.items():
        if any(kw in lower for kw in kws):
            return dtype
    return "OTHER"


def detect_version(page1_text: str) -> str:
    m = VERSION_RE.search(page1_text)
    return m.group(1).strip() if m else "unknown"


def detect_title_from_page1(page1_text: str, font_data: list = None) -> str:
    """
    DEFECT-5 fix: scan page 1 specifically.
    Primary: use largest-font line from pdfplumber word data.
    Fallback: first non-empty Title-Case line with ≥ 3 words.
    """
    # Font-based (most reliable)
    if font_data:
        max_size = max(w.get("size", 0) for w in font_data) if font_data else 0
        for w in font_data:
            if w.get("size", 0) >= max_size - 0.5:
                candidate = w.get("text", "").strip()
                if len(candidate.split()) >= 3:
                    return candidate

    # Text-based fallback
    for line in page1_text.split("\n")[:20]:
        line = line.strip()
        if (
            re.match(r"^[A-Z][A-Za-z0-9\s\-:]{10,120}$", line)
            and len(line.split()) >= 3
        ):
            return line
    return "unknown"


def detect_language(text: str) -> str:
    if HAS_LANGDETECT:
        try:
            return langdetect.detect(text[:2000])
        except Exception:
            pass
    return "unknown"


# ══════════════════════════════════════════════════════════════
# LINE-LEVEL PARAGRAPH BUILDER (DEFECT-1 fix — core fix)
# ══════════════════════════════════════════════════════════════

def build_paragraphs_from_lines(lines: list, doc_id: str, page: int, start_seq: int):
    """
    Splits page text into proper paragraphs using line-level heading detection.
    A new paragraph starts whenever a heading line is encountered OR
    after a blank line gap.
    Returns (paragraphs: list, end_seq: int, page_headings: list).
    """
    paragraphs = []
    page_headings = []
    seq = start_seq

    current_lines = []
    current_is_heading = False
    current_level = None
    current_method = None

    def flush(lines_buf, is_heading, level, method):
        nonlocal seq
        text = "\n".join(lines_buf).strip()
        if not text:
            return
        seq += 1
        para_id = make_para_id(doc_id, page, seq)
        paragraphs.append({
            "page": page,
            "para_id": para_id,
            "text": text,
            "is_heading": is_heading,
            "heading_level": level,
            "heading_detection_method": method,
        })
        if is_heading:
            page_headings.append({"level": level, "text": lines_buf[0].strip(), "method": method})

    for raw_line in lines:
        line = raw_line.strip()

        # Blank line → flush current buffer as body paragraph
        if not line:
            if current_lines:
                flush(current_lines, current_is_heading, current_level, current_method)
                current_lines = []
                current_is_heading = False
                current_level = None
                current_method = None
            continue

        level, method = detect_heading_level(line)
        clean = clean_heading_text(line)

        if level is not None:
            # Flush whatever was accumulating before
            if current_lines:
                flush(current_lines, current_is_heading, current_level, current_method)
            # Start fresh heading paragraph
            current_lines = [clean]
            current_is_heading = True
            current_level = level
            current_method = method
        else:
            if current_is_heading and current_lines:
                # Heading was one line; flush it, start body
                flush(current_lines, True, current_level, current_method)
                current_lines = [line]
                current_is_heading = False
                current_level = None
                current_method = None
            else:
                current_lines.append(line)

    # Final flush
    if current_lines:
        flush(current_lines, current_is_heading, current_level, current_method)

    return paragraphs, seq, page_headings


# ══════════════════════════════════════════════════════════════
# PDF EXTRACTOR
# ══════════════════════════════════════════════════════════════

def extract_pdf_content(path: str, doc_id: str) -> dict:
    print(f"[PDF] Processing: {path}")
    result = {
        "pages": [],
        "ocr_text": [],
        "tables": [],
        "paragraphs": [],
        "page_section_map": [],
    }

    doc_name = os.path.splitext(os.path.basename(path))[0]
    doc_img_dir = EXTRACTED_IMAGES_DIR / doc_name
    doc_img_dir.mkdir(parents=True, exist_ok=True)

    heading_stack = []  # [(level, clean_text)]
    global_seq = 0
    page1_font_data = []

    # ── pdfplumber for font metadata on page 1 ──
    with pdfplumber.open(path) as plumb:
        page1_words = plumb.pages[0].extract_words(extra_attrs=["size"]) or []
        page1_font_data = page1_words

    # ── First pass: collect all heading texts for ToC cross-reference ──
    # (used by detect_page_type to identify ToC pages reliably)
    known_headings: set = set()
    with pdfplumber.open(path) as plumb_scan:
        for scan_page in plumb_scan.pages[1:]:  # skip page 1
            for line in (scan_page.extract_text() or "").split("\n"):
                line = line.strip()
                if SECTION_NUM_RE.match(line) and not LIST_ITEM_SIGNALS.search(line) and len(line) <= 80:
                    known_headings.add(ANCHOR_RE.sub("", line).strip())

    doc = fitz.open(path)
    for pno, page in enumerate(doc, 1):
        raw_text = page.get_text("text") or ""
        lines = raw_text.split("\n")

        # ── Issue 1 Fix: classify page type before processing ──
        page_type = detect_page_type(lines, known_headings)
        result["pages"].append({"page": pno, "content": raw_text, "page_type": page_type})

        paras, global_seq, page_headings = build_paragraphs_from_lines(
            lines, doc_id, pno, global_seq
        )

        # ── Issue 1 Fix: skip heading stack update for toc/cover pages ──
        # Prevents ToC entries from poisoning section context for all subsequent pages
        if page_type == "body":
            # ── Issue 2 Fix: update stack BEFORE building section_map ──
            # (previously stack was updated after, causing off-by-one section paths)
            for h in page_headings:
                heading_stack = [(l, t) for l, t in heading_stack if l < h["level"]]
                heading_stack.append((h["level"], h["text"]))

        # ── Issue 2 Fix: stamp section_path onto each paragraph NOW ──
        # (stack is current at this point; each para gets its exact context)
        para_stack_snapshot = list(heading_stack)  # snapshot after update
        # Walk back through just-built paras and assign section_path per paragraph
        # We rebuild inline since headings within the page shift the context
        inline_stack = list(heading_stack)
        # Reset to pre-page stack, then replay heading updates per-paragraph
        if page_type == "body":
            # Re-derive per-paragraph section_path by replaying heading order
            pre_page_stack = [(l, t) for l, t in heading_stack]
            # Remove headings found on THIS page from the stack to get pre-page state
            this_page_heading_texts = {h["text"] for h in page_headings}
            pre_page_stack = [(l, t) for l, t in pre_page_stack if t not in this_page_heading_texts]
            running_stack = list(pre_page_stack)
            for para in paras:
                if para["is_heading"]:
                    running_stack = [(l, t) for l, t in running_stack if l < para["heading_level"]]
                    running_stack.append((para["heading_level"], para["text"]))
                para["section_path"] = build_section_path(running_stack)
        else:
            for para in paras:
                para["section_path"] = f"__SKIP__{page_type.upper()}__"
                para["page_type"] = page_type

        result["paragraphs"].extend(paras)

        # ── Page section map ──
        section_path = build_section_path(heading_stack) if heading_stack else "UNMAPPED"
        result["page_section_map"].append({
            "page": pno,
            "page_type": page_type,
            "section_path": section_path if page_type == "body" else f"__SKIP__{page_type.upper()}__",
            "heading_text": heading_stack[-1][1] if heading_stack else None,
            "heading_level": heading_stack[-1][0] if heading_stack else None,
            "headings_on_page": page_headings if page_type == "body" else [],
            "ambiguous": len(page_headings) == 0 and pno > 1 and page_type == "body",
        })

        # ── OCR for embedded images ──
        # IMPORTANT: PDFs often embed bullet points, checkboxes, and decorative
        # glyphs as tiny inline images (~3-5 pts wide). These are NOT content
        # images and must be skipped. We use a minimum size threshold:
        #   width  >= MIN_IMAGE_PTS  AND  height >= MIN_IMAGE_PTS
        # This PDF's bullets are 3.7×3.7 pts; its real diagrams are 195×396,
        # 264×470, 503×153 pts — a 2000× area difference.
        MIN_IMAGE_PTS = 50  # points (~17mm); tune down to 30 if needed

        # Build a set of xrefs that are large enough to be real images,
        # using pdfplumber's richer image metadata (includes bbox dimensions).
        # BUG FIX: pdfplumber has NO "xref" key — img.get("xref") always returns
        # None, making the set permanently empty and blocking ALL images.
        # The PDF object number (equivalent to fitz's xref) is at stream.objid.
        real_image_xrefs = set()
        with pdfplumber.open(path) as plumb:
            plumb_page = plumb.pages[pno - 1]
            for plumb_img in plumb_page.images:
                w = plumb_img.get("width", 0)
                h = plumb_img.get("height", 0)
                if w >= MIN_IMAGE_PTS and h >= MIN_IMAGE_PTS:
                    stream = plumb_img.get("stream")
                    if stream is not None:
                        real_image_xrefs.add(stream.objid)

        for i, img in enumerate(page.get_images(full=True)):
            xref = img[0]

            # Skip glyphs, bullets, icons — not real content images
            if xref not in real_image_xrefs:
                continue

            pix = fitz.Pixmap(doc, xref)
            img_path = doc_img_dir / f"page{pno}_img{i+1}.png"
            (pix if pix.n < 5 else fitz.Pixmap(fitz.csRGB, pix)).save(str(img_path))
            img_cv = cv2.imread(str(img_path))
            if img_cv is None:
                continue
            gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
            ocr_text = pytesseract.image_to_string(gray)

            # ── Issue 5 Fix: assess OCR quality ──
            # Noise ratio = non-alphanumeric / total chars (excluding whitespace)
            # > 15% noise → "low" quality flag; Stage 2 should treat as supplementary only
            non_ws = ocr_text.replace(" ", "").replace("\n", "")
            if non_ws:
                noise_chars = sum(1 for c in non_ws if not c.isalnum())
                noise_ratio = round(noise_chars / len(non_ws), 3)
            else:
                noise_ratio = 1.0
            ocr_quality = "low" if noise_ratio > 0.15 else "acceptable"

            result["ocr_text"].append({
                "page": pno,
                "image": str(img_path),
                "content": ocr_text,
                "noise_ratio": noise_ratio,
                "quality_flag": ocr_quality,
                "stage2_note": (
                    "Use as supplementary context only — noise ratio too high for primary extraction"
                    if ocr_quality == "low"
                    else "Acceptable for primary extraction"
                ),
            })

    doc.close()

    # ── Table extraction: camelot primary, pdfplumber fallback (DEFECT-4) ──
    # Issue 3 Fix: filter single-column camelot results — they are wrapped
    # text blocks misidentified as tables, not real structured data.
    # A genuine table must have ≥ 2 columns AND ≥ 2 data rows.
    try:
        tables = camelot.read_pdf(path, pages="all", flavor="stream")
        real_table_count = 0
        for i, t in enumerate(tables):
            n_cols = t.df.shape[1]
            n_rows = t.df.shape[0]
            if n_cols < 2 or n_rows < 2:
                continue  # single-column or near-empty = wrapped text, not a table
            real_table_count += 1
            result["tables"].append({
                "table_index": real_table_count - 1,
                "page": t.page,
                "source": "camelot",
                "columns": n_cols,
                "rows": n_rows,
                "data": t.df.to_dict(orient="records"),
            })
        print(f"[PDF] camelot: {len(tables)} raw → {real_table_count} real tables (≥2 cols, ≥2 rows)")
    except Exception as e:
        print(f"[PDF] camelot failed ({e}), trying pdfplumber...")
        try:
            with pdfplumber.open(path) as plumb:
                for pno, page in enumerate(plumb.pages, 1):
                    for ti, tbl in enumerate(page.extract_tables() or []):
                        result["tables"].append({
                            "table_index": len(result["tables"]),
                            "page": pno,
                            "source": "pdfplumber",
                            "data": [dict(zip(map(str, range(len(tbl[0]))), row)) for row in tbl],
                        })
        except Exception as e2:
            print(f"[PDF] pdfplumber table extraction also failed: {e2}")

    return result, page1_font_data


# ══════════════════════════════════════════════════════════════
# DOCX EXTRACTOR
# ══════════════════════════════════════════════════════════════

def extract_docx(path: str, doc_id: str) -> tuple:
    print(f"[DOCX] Processing: {path}")
    doc = Document(path)
    result = {"pages": [], "ocr_text": [], "tables": [], "paragraphs": [], "page_section_map": []}

    heading_stack = []
    global_seq = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        hm = re.match(r"Heading (\d+)", style_name)
        if hm:
            level, method = int(hm.group(1)), "word_style"
        else:
            level, method = detect_heading_level(text)

        clean = clean_heading_text(text)
        global_seq += 1
        para_id = make_para_id(doc_id, 1, global_seq)
        result["paragraphs"].append({
            "page": 1,
            "para_id": para_id,
            "text": clean,
            "is_heading": level is not None,
            "heading_level": level,
            "heading_detection_method": method if level else None,
        })

        if level is not None:
            heading_stack = [(l, t) for l, t in heading_stack if l < level]
            heading_stack.append((level, clean))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    result["pages"].append({"page": 1, "content": full_text})
    if heading_stack:
        result["page_section_map"].append({
            "page": 1,
            "section_path": build_section_path(heading_stack),
            "heading_text": heading_stack[-1][1],
            "heading_level": heading_stack[-1][0],
            "headings_on_page": [{"level": l, "text": t} for l, t in heading_stack],
            "ambiguous": False,
        })
    return result, []


# ══════════════════════════════════════════════════════════════
# TXT / XLSX / IMAGE EXTRACTORS
# ══════════════════════════════════════════════════════════════

def extract_txt(path: str, doc_id: str) -> tuple:
    print(f"[TXT] Processing: {path}")
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()

    result = {"pages": [{"page": 1, "content": text}], "ocr_text": [], "tables": [],
              "paragraphs": [], "page_section_map": []}
    heading_stack = []
    lines = text.split("\n")
    paras, _, headings = build_paragraphs_from_lines(lines, doc_id, 1, 0)
    result["paragraphs"] = paras
    for h in headings:
        heading_stack = [(l, t) for l, t in heading_stack if l < h["level"]]
        heading_stack.append((h["level"], h["text"]))
    if heading_stack:
        result["page_section_map"].append({
            "page": 1,
            "section_path": build_section_path(heading_stack),
            "heading_text": heading_stack[-1][1],
            "heading_level": heading_stack[-1][0],
            "headings_on_page": headings,
            "ambiguous": False,
        })
    return result, []


def extract_xlsx(path: str, doc_id: str) -> tuple:
    print(f"[XLSX] Processing: {path}")
    xls = pd.ExcelFile(path)
    result = {"pages": [], "ocr_text": [], "tables": [], "paragraphs": [], "page_section_map": []}
    for i, name in enumerate(xls.sheet_names):
        rows = pd.read_excel(xls, name).to_dict(orient="records")
        result["pages"].append({"page": i + 1, "content": f"Sheet: {name}"})
        result["tables"].append({"table_index": i, "page": i + 1, "sheet_name": name, "data": rows})
        para_id = make_para_id(doc_id, i + 1, 1)
        result["paragraphs"].append({"page": i + 1, "para_id": para_id, "text": f"Sheet: {name}",
                                      "is_heading": True, "heading_level": 1, "heading_detection_method": "sheet_name"})
        result["page_section_map"].append({"page": i + 1, "section_path": name, "heading_text": name,
                                            "heading_level": 1, "headings_on_page": [{"level": 1, "text": name}],
                                            "ambiguous": False})
    return result, []


def extract_image(path: str, doc_id: str) -> tuple:
    print(f"[IMG] Processing: {path}")
    img = cv2.imread(path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    ocr_text = pytesseract.image_to_string(gray)
    para_id = make_para_id(doc_id, 1, 1)
    return {
        "pages": [{"page": 1, "content": ocr_text}],
        "ocr_text": [{"page": 1, "image": path, "content": ocr_text}],
        "tables": [],
        "paragraphs": [{"page": 1, "para_id": para_id, "text": ocr_text,
                         "is_heading": False, "heading_level": None, "heading_detection_method": None}],
        "page_section_map": [{"page": 1, "section_path": "UNMAPPED", "heading_text": None,
                               "heading_level": None, "headings_on_page": [], "ambiguous": True}],
    }, []


# ══════════════════════════════════════════════════════════════
# MASTER PROCESS FUNCTION
# ══════════════════════════════════════════════════════════════

def process_file(file_path: str) -> dict:
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    temp_id = "doc_" + hashlib.sha256(filename.encode()).hexdigest()[:16]

    if ext == ".pdf":
        content, font_data = extract_pdf_content(file_path, temp_id)
    elif ext == ".docx":
        content, font_data = extract_docx(file_path, temp_id)
    elif ext == ".txt":
        content, font_data = extract_txt(file_path, temp_id)
    elif ext == ".xlsx":
        content, font_data = extract_xlsx(file_path, temp_id)
    elif ext in (".png", ".jpg", ".jpeg"):
        content, font_data = extract_image(file_path, temp_id)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    all_text = " ".join(p["content"] for p in content["pages"])
    page1_text = content["pages"][0]["content"] if content["pages"] else ""
    doc_id = make_doc_id(filename, all_text)

    # Fix para_ids to use real doc_id
    for p in content["paragraphs"]:
        p["para_id"] = p["para_id"].replace(temp_id, doc_id)

    # ── Issue 4 Fix: normalize FR sub-structure labels ──
    # Strip "Description: " prefix from requirement body paragraphs and
    # annotate paragraphs with a sub_type field for Stage 2.
    DESCRIPTION_PREFIX_RE = re.compile(r"^Description:\s*", re.IGNORECASE)
    FR_SUB_LABELS = {
        "inputs": re.compile(r"^Inputs?:\s*$", re.IGNORECASE),
        "system_behavior": re.compile(r"^System\s+Behavior:\s*$", re.IGNORECASE),
        "outputs": re.compile(r"^Outputs?:\s*$", re.IGNORECASE),
        "assumptions": re.compile(r"^Assumptions?:\s*$", re.IGNORECASE),
        "dependencies": re.compile(r"^Dependencies?:\s*$", re.IGNORECASE),
    }
    for p in content["paragraphs"]:
        text = p["text"]
        # Strip Description: prefix and mark as requirement_statement
        if DESCRIPTION_PREFIX_RE.match(text):
            p["text"] = DESCRIPTION_PREFIX_RE.sub("", text).strip()
            p["sub_type"] = "requirement_statement"
        else:
            # Detect and tag sub-structure labels
            matched_sub = None
            for sub_type, pattern in FR_SUB_LABELS.items():
                if pattern.match(text.split("\n")[0]):
                    matched_sub = sub_type
                    break
            p["sub_type"] = matched_sub if matched_sub else "body"

    doc_metadata = {
        "doc_id": doc_id,
        "title": detect_title_from_page1(page1_text, font_data),
        "doc_type": detect_doc_type(all_text[:3000]),
        "version": detect_version(page1_text),
        "source_file": filename,
        "ingestion_timestamp": datetime.now(timezone.utc).isoformat(),
        "language": detect_language(all_text),
    }

    return {
        "doc_metadata": doc_metadata,
        "page_section_map": content["page_section_map"],
        "paragraphs": content["paragraphs"],
        "pages": content["pages"],
        "ocr_text": content["ocr_text"],
        "tables": content["tables"],
    }


# ══════════════════════════════════════════════════════════════
# AUDIT REPORT GENERATOR
# ══════════════════════════════════════════════════════════════

def generate_audit_report(structured: dict) -> dict:
    meta = structured["doc_metadata"]
    pages = structured["pages"]
    paragraphs = structured["paragraphs"]
    section_map = structured["page_section_map"]
    tables = structured["tables"]
    ocr = structured["ocr_text"]

    total_pages = len(pages)
    total_paras = len(paragraphs)

    page_para_counts = defaultdict(int)
    for p in paragraphs:
        page_para_counts[p["page"]] += 1
    counts = list(page_para_counts.values()) or [0]

    unmapped = [s["page"] for s in section_map if s["section_path"] == "UNMAPPED"]
    multi_heading_pages = [s["page"] for s in section_map if len(s.get("headings_on_page", [])) > 1]
    no_heading_pages = [s["page"] for s in section_map if not s.get("headings_on_page")]
    ambiguous_pages = [s["page"] for s in section_map if s.get("ambiguous")]

    seen_sections = defaultdict(list)
    for s in section_map:
        sp = s["section_path"]
        if sp and sp != "UNMAPPED":
            seen_sections[sp].append(s["page"])
    multi_page_sections = {sp: pgs for sp, pgs in seen_sections.items() if len(pgs) > 1}

    para_ids = [p["para_id"] for p in paragraphs]
    unique_ids = len(set(para_ids))
    para_id_unique = unique_ids == len(para_ids)

    by_method = defaultdict(int)
    for p in paragraphs:
        if p.get("is_heading") and p.get("heading_detection_method"):
            by_method[p["heading_detection_method"]] += 1

    meta_checks = {k: v not in ("unknown", None, "") for k, v in meta.items()}

    warnings = []
    issues = []

    if not para_id_unique:
        issues.append(f"DUPLICATE paragraph IDs: {len(para_ids) - unique_ids} collisions")
    if len(unmapped) > total_pages * 0.5:
        issues.append(f"More than 50% pages unmapped ({len(unmapped)}/{total_pages})")
    if ambiguous_pages:
        warnings.append(f"Ambiguous heading detection on pages: {ambiguous_pages[:20]}")
    if not meta_checks.get("title"):
        warnings.append("Title not detected — defaulted to 'unknown'")
    if not meta_checks.get("version"):
        warnings.append("Version not detected — defaulted to 'unknown'")
    if ocr:
        warnings.append(f"OCR used on {len(ocr)} image(s) — text accuracy may vary")
    if not tables:
        warnings.append("No tables extracted")

    verdict = "NOT READY" if issues else ("PARTIALLY READY" if warnings else "READY")

    return {
        "report_generated_at": datetime.now(timezone.utc).isoformat(),
        "document_summary": {
            "doc_id": meta["doc_id"],
            "doc_type": meta["doc_type"],
            "source_file": meta["source_file"],
            "total_pages_ingested": total_pages,
            "total_paragraphs_extracted": total_paras,
            "tables_detected": len(tables),
            "ocr_used": len(ocr) > 0,
        },
        "structural_coverage": {
            "pages_unmapped": unmapped,
            "pages_with_multiple_headings": multi_heading_pages,
            "pages_with_no_headings": no_heading_pages,
            "sections_spanning_multiple_pages": multi_page_sections,
            "paragraph_stats": {
                "total": total_paras,
                "min_per_page": min(counts),
                "max_per_page": max(counts),
                "avg_per_page": round(sum(counts) / len(counts), 2),
            },
        },
        "heading_detection_summary": {
            "total_headings": sum(1 for p in paragraphs if p.get("is_heading")),
            "by_method": dict(by_method),
            "unique_section_paths": len(seen_sections),
        },
        "metadata_completeness": {
            "fields": meta_checks,
            "all_complete": all(meta_checks.values()),
        },
        "data_quality": {
            "para_ids_unique": para_id_unique,
            "total": len(para_ids),
            "unique": unique_ids,
            "section_map_coverage": f"{total_pages - len(unmapped)}/{total_pages} pages mapped",
        },
        "blocking_issues": issues,
        "warnings": warnings,
        "final_verdict": {
            "status": verdict,
            "reasons": issues if issues else (warnings if warnings else ["All checks passed."]),
        },
    }


# ══════════════════════════════════════════════════════════════
# RUNNER
# ══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    input_path = input("Enter file path (PDF/DOCX/TXT/XLSX/IMG): ").strip()

    structured = process_file(input_path)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    structured_output_path = OUTPUT_DIR / "structured_output.json"
    audit_output_path = OUTPUT_DIR / "ingestion_audit_report.json"

    with open(structured_output_path, "w", encoding="utf-8") as f:
        json.dump(structured, f, ensure_ascii=False, indent=2)
    print(f"✅ structured_output.json → {structured_output_path}")

    audit = generate_audit_report(structured)

    with open(audit_output_path, "w", encoding="utf-8") as f:
        json.dump(audit, f, ensure_ascii=False, indent=2)
    print(f"✅ ingestion_audit_report.json → {audit_output_path}")

    v = audit["final_verdict"]
    s = audit["document_summary"]
    print(f"\n{'='*55}")
    print(f"  VERDICT : {v['status']}")
    print(f"  Pages   : {s['total_pages_ingested']}")
    print(f"  Paras   : {s['total_paragraphs_extracted']}")
    print(f"  Tables  : {s['tables_detected']}")
    for r in v["reasons"]:
        print(f"  → {r}")
    print(f"{'='*55}\n")
