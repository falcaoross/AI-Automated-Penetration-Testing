"""
Module 5.1: Acceptance Story Ingestion & Structuring

This module extracts structured acceptance units from raw User Acceptance Story (UAS) documents.
It performs STRUCTURAL extraction only - no semantic validation, no CRU comparison, no gap detection.

Core Responsibility:
    Convert raw acceptance documents (PDF/DOCX/TXT) into traceable, structured acceptance units
    that are ready for downstream semantic comparison modules.

Constraints:
    - Preserves exact text (no paraphrasing)
    - No semantic interpretation
    - No validation against CRUs
    - Over-inclusion preferred (downstream modules will filter)
    - Traceability is paramount
"""

import json
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

# PDF processing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX processing
try:
    from docx import Document
except ImportError:
    Document = None


class AcceptanceUnit:
    """Represents a single extracted acceptance expectation."""
    
    def __init__(
        self,
        uas_id: str,
        text: str,
        source_section: Optional[str],
        source_type: str,
        page: Optional[int],
        paragraph: Optional[int],
        source_file: str,
        scenario_id: Optional[str] = None,
        is_header: bool = False,
        actor_role: Optional[str] = None
    ):
        self.uas_id = uas_id
        self.text = text
        self.source_section = source_section
        self.source_type = source_type
        self.page = page
        self.paragraph = paragraph
        self.source_file = source_file
        self.scenario_id = scenario_id
        self.is_header = is_header
        self.actor_role = actor_role
    
    def to_dict(self) -> Dict:
        """Convert to output JSON schema."""
        return {
            "uas_id": self.uas_id,
            "text": self.text,
            "source_section": self.source_section,
            "source_type": self.source_type,
            "scenario_id": self.scenario_id,
            "is_header": self.is_header,
            "actor_role": self.actor_role,
            "traceability": {
                "page": self.page,
                "paragraph": self.paragraph,
                "source_file": self.source_file
            }
        }


class AcceptanceStoryIngester:
    """Main ingestion engine for acceptance documents."""
    
    # Structural cue phrases for segmentation
    ACCEPTANCE_CUES = [
        r'\bGiven\b',
        r'\bWhen\b',
        r'\bThen\b',
        r'\bAnd\b',
        r'\bScenario:',
        r'\bAcceptance Story\b',
        r'\bAcceptance Criteria\b',
        r'\breturns HTTP',
        r'\bdisplays?\b',
        r'\bredirected to\b',
        r'\bshall\b',
        r'\bmust\b',
        r'\bshould\b',
        r'\bexpect',
        r'\bwant to\b',
    ]
    
    # Source type classification keywords
    TYPE_MARKERS = {
        'user_story': [
            r'As an?\b',
            r'I want\b',
            r'so that\b',
            r'^User Story$',
        ],
        'acceptance_story': [
            r'Given\b',
            r'When\b',
            r'Then\b',
            r'Scenario:',
            r'Acceptance Story',
        ],
        'example': [
            r'Example:',
            r'For instance',
            r'\be\.g\.',
            r'For example',
        ],
        'note': [
            r'Note:',
            r'Important:',
            r'Reminder:',
        ]
    }
    
    def __init__(self, source_file: str):
        self.source_file = Path(source_file).name
        self.source_path = source_file
        self.units: List[AcceptanceUnit] = []
        self.unit_counter = 0
        self.pages_processed = 0
        self.scenario_counter = 0
        self.current_scenario_id: Optional[str] = None
        
    def extract_from_pdf(self) -> List[Tuple[int, str, Optional[str]]]:
        """Extract text from PDF with page numbers and sections."""
        if pdfplumber is None:
            raise ImportError("pdfplumber not installed")
        
        sections = []
        current_section = None
        
        with pdfplumber.open(self.source_path) as pdf:
            self.pages_processed = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if not text:
                    continue
                
                section_match = re.search(r'^[🔹●■▪]\s*(?:ROLE:\s*)?([A-Z\s]+)$', text, re.MULTILINE)
                if section_match:
                    current_section = section_match.group(1).strip()
                
                sections.append((page_num, text, current_section))
        
        return sections
    
    def extract_from_docx(self) -> List[Tuple[int, str, Optional[str]]]:
        """Extract text from DOCX with paragraph tracking and sections."""
        if Document is None:
            raise ImportError("python-docx not installed")
        
        doc = Document(self.source_path)
        sections = []
        current_section = None
        
        for para_idx, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if not text:
                continue
            
            if para.style.name.startswith('Heading') or text.isupper():
                current_section = text
            
            sections.append((para_idx, text, current_section))
        
        self.pages_processed = len(doc.paragraphs)
        return sections
    
    def extract_from_txt(self) -> List[Tuple[int, str, Optional[str]]]:
        """Extract text from plain text file."""
        with open(self.source_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        self.pages_processed = len(text.split('\n'))
        
        # For TXT files, return the entire content as one block
        # (unlike PDF where we process page-by-page)
        return [(1, text, None)]
    
    def segment_text(self, text: str, page_num: int, section: Optional[str]) -> List[Dict]:
        """Segment text into candidate acceptance units."""
        candidates = []
        lines = text.split('\n')
        
        current_user_story = None
        current_acceptance_story = None
        paragraph_counter = 0
        
        # Buffer for multi-line user stories
        user_story_buffer = []
        user_story_start_para = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            paragraph_counter += 1
            
            # User Story headers
            if re.match(r'^(US [A-Z]+ \d+|User Story\s*$)', line, re.IGNORECASE):
                if user_story_buffer:
                    merged_text = ' '.join(user_story_buffer)
                    if self._is_valid_unit(merged_text):
                        candidates.append({
                            'text': merged_text,
                            'page': page_num,
                            'paragraph': user_story_start_para,
                            'section': section,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                    user_story_buffer = []
                    user_story_start_para = None
                
                current_user_story = line
                current_acceptance_story = None
                self.current_scenario_id = None
                continue
            
            # Acceptance Story headers
            if re.match(r'^Acceptance Story\s*$', line, re.IGNORECASE):
                if user_story_buffer:
                    merged_text = ' '.join(user_story_buffer)
                    if self._is_valid_unit(merged_text):
                        candidates.append({
                            'text': merged_text,
                            'page': page_num,
                            'paragraph': user_story_start_para,
                            'section': section,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                    user_story_buffer = []
                    user_story_start_para = None
                
                current_acceptance_story = line
                continue
            
            # Scenario headers
            if self._is_scenario_header(line):
                if user_story_buffer:
                    merged_text = ' '.join(user_story_buffer)
                    if self._is_valid_unit(merged_text):
                        candidates.append({
                            'text': merged_text,
                            'page': page_num,
                            'paragraph': user_story_start_para,
                            'section': section,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                    user_story_buffer = []
                    user_story_start_para = None
                
                self.scenario_counter += 1
                next_scenario_id = f"SCN_{self.scenario_counter:03d}"
                
                candidates.append({
                    'text': line,
                    'page': page_num,
                    'paragraph': paragraph_counter,
                    'section': section or current_user_story,
                    'scenario_id': next_scenario_id,
                    'is_header': True,
                })
                
                self.current_scenario_id = next_scenario_id
                continue
            
            # Start of user story "As a..."
            if re.match(r'^As an?\s+', line, re.IGNORECASE):
                if user_story_buffer:
                    merged_text = ' '.join(user_story_buffer)
                    if self._is_valid_unit(merged_text):
                        candidates.append({
                            'text': merged_text,
                            'page': page_num,
                            'paragraph': user_story_start_para,
                            'section': section,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                
                user_story_buffer = [line]
                user_story_start_para = paragraph_counter
                continue
            
            # Continue buffering user story lines
            if user_story_buffer:
                # Check if this is a structural boundary
                is_structural = (
                    re.match(r'^(Given|When|Then|And)\b', line, re.IGNORECASE) or
                    re.match(r'^[\-•●▪*]\s+', line) or
                    re.match(r'^\d+\.\s+', line) or
                    re.match(r'^Scenario', line, re.IGNORECASE) or
                    re.match(r'^Acceptance Story\s*$', line, re.IGNORECASE) or
                    re.match(r'^(US [A-Z]+ \d+|User Story\s*$)', line, re.IGNORECASE)
                )
                
                if not is_structural:
                    user_story_buffer.append(line)
                    
                    # Check if pattern is complete
                    full_text = ' '.join(user_story_buffer)
                    has_complete_pattern = (
                        re.search(r'\b(I want|I should|I need)\b', full_text, re.IGNORECASE) and 
                        re.search(r'\b(so that|in order to)\b', full_text, re.IGNORECASE)
                    )
                    
                    # Flush if complete (period preferred but not required)
                    if has_complete_pattern and (line.endswith('.') or len(user_story_buffer) >= 3):
                        if self._is_valid_unit(full_text):
                            candidates.append({
                                'text': full_text,
                                'page': page_num,
                                'paragraph': user_story_start_para,
                                'section': section,
                                'scenario_id': self.current_scenario_id,
                                'is_header': False,
                            })
                        user_story_buffer = []
                        user_story_start_para = None
                    continue
                else:
                    # Structural boundary hit - flush buffer
                    merged_text = ' '.join(user_story_buffer)
                    
                    # Only flush if it's a valid, complete user story
                    # Otherwise, the buffer might be incomplete (e.g., buffered before structural keyword was seen)
                    has_complete_pattern = (
                        re.search(r'\b(I want|I should|I need)\b', merged_text, re.IGNORECASE) and 
                        re.search(r'\b(so that|in order to)\b', merged_text, re.IGNORECASE)
                    )
                    
                    if has_complete_pattern and self._is_valid_unit(merged_text):
                        candidates.append({
                            'text': merged_text,
                            'page': page_num,
                            'paragraph': user_story_start_para,
                            'section': section,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                    user_story_buffer = []
                    user_story_start_para = None
                    # Fall through to process current line
            
            # Bullet points / numbered lists
            if re.match(r'^[\-•●▪*]\s+', line) or re.match(r'^\d+\.\s+', line):
                cleaned = re.sub(r'^[\-•●▪*\d.]+\s+', '', line)
                if self._is_valid_unit(cleaned):
                    candidates.append({
                        'text': cleaned,
                        'page': page_num,
                        'paragraph': paragraph_counter,
                        'section': section or current_user_story,
                        'scenario_id': self.current_scenario_id,
                        'is_header': False,
                    })
                continue
            
            # Gherkin steps
            if re.match(r'^(Given|When|Then|And)\b', line, re.IGNORECASE):
                if self._is_valid_unit(line):
                    candidates.append({
                        'text': line,
                        'page': page_num,
                        'paragraph': paragraph_counter,
                        'section': current_acceptance_story or current_user_story or section,
                        'scenario_id': self.current_scenario_id,
                        'is_header': False,
                    })
                continue
            
            # Narrative acceptance expectations
            if any(re.search(pattern, line, re.IGNORECASE) for pattern in self.ACCEPTANCE_CUES):
                sentences = re.split(r'(?<=[.!?])\s+', line)
                for sentence in sentences:
                    if self._is_valid_unit(sentence):
                        candidates.append({
                            'text': sentence,
                            'page': page_num,
                            'paragraph': paragraph_counter,
                            'section': section or current_user_story,
                            'scenario_id': self.current_scenario_id,
                            'is_header': False,
                        })
                continue
        
        # Flush remaining buffer
        if user_story_buffer:
            merged_text = ' '.join(user_story_buffer)
            if self._is_valid_unit(merged_text):
                candidates.append({
                    'text': merged_text,
                    'page': page_num,
                    'paragraph': user_story_start_para,
                    'section': section,
                    'scenario_id': self.current_scenario_id,
                    'is_header': False,
                })
        
        return candidates
    
    def _is_valid_unit(self, text: str) -> bool:
        """Validate if text qualifies as an acceptance unit."""
        if not text or len(text.strip()) < 6:
            return False
        
        words = text.split()
        if len(words) < 6:
            return False
        
        has_expectation = any([
            re.search(r'\b(shall|must|should|will|can|may|want|expect|need)\b', text, re.IGNORECASE),
            re.search(r'\b(click|enter|submit|redirect|display|return|create|update|delete|validate|filter|handle)\b', text, re.IGNORECASE),
            re.search(r'\b(Given|When|Then|And)\b', text, re.IGNORECASE),
        ])
        
        if not has_expectation:
            return False
        
        if re.match(r'^[\d\s\-_=.]+$', text):
            return False
        
        if re.match(r'^(Page\s+\d+|Section\s+\d+)', text, re.IGNORECASE):
            return False
        
        return True
    
    def _is_scenario_header(self, text: str) -> bool:
        """Check if text is a scenario header."""
        return bool(re.match(r'^(Scenario|Scenario Outline|Acceptance Scenario):', text, re.IGNORECASE))
    
    def _extract_actor_role(self, text: str, is_header: bool) -> Optional[str]:
        """Extract actor role using strict priority logic."""
        # Rule 1: Headers never have actor roles
        if is_header:
            return None
        
        text_lower = text.lower()
        text_stripped = text.strip()
        
        # Rule 2: Text starts with "the system"
        if text_lower.startswith('the system'):
            return 'system'
        
        # Rule 3: Contains "super admin" or "admin"
        if 'super admin' in text_lower:
            return 'admin'
        if 'admin' in text_lower:
            return 'admin'
        
        # Rule 4: Text starts with "as a" - extract role from phrase
        as_a_match = re.match(r'^as an?\s+([^,]+)', text_lower)
        if as_a_match:
            role_phrase = as_a_match.group(1).strip()
            if 'admin' in role_phrase or 'administrator' in role_phrase:
                return 'admin'
            if 'user' in role_phrase:
                return 'user'
            if 'system' in role_phrase:
                return 'system'
            if 'external' in role_phrase or 'third-party' in role_phrase or 'third party' in role_phrase:
                return 'external'
        
        # Rule 5: Gherkin steps - only assign if subject is explicitly system/admin
        if re.match(r'^(given|when|then|and)\b', text_lower):
            # Extract subject after the keyword
            subject_match = re.match(r'^(given|when|then|and)\s+the\s+(\w+)', text_lower)
            if subject_match:
                subject = subject_match.group(2)
                if subject == 'system':
                    return 'system'
                if subject == 'admin' or subject == 'administrator':
                    return 'admin'
            
            # Check if "the user" is the explicit subject
            if re.match(r'^(given|when|then|and)\s+the\s+user\b', text_lower):
                return 'user'
            
            # Otherwise return None for Gherkin steps
            return None
        
        # Rule 6: Check for external systems
        if 'external system' in text_lower or 'third-party' in text_lower or 'third party' in text_lower:
            return 'external'
        
        # If unsure, return None
        return None
    
    def classify_source_type(self, text: str) -> str:
        """Classify unit source type using strict keyword rules."""
        # Strict rules
        if re.match(r'^(Given|When|Then|And)\b', text, re.IGNORECASE):
            return 'acceptance_story'
        
        if re.match(r'^Scenario', text, re.IGNORECASE):
            return 'acceptance_story'
        
        if re.match(r'^As an?\b', text, re.IGNORECASE):
            return 'user_story'
        
        if re.match(r'^(Example:|For instance|e\.g\.)', text, re.IGNORECASE):
            return 'example'
        
        # Fallback to TYPE_MARKERS
        for type_name, patterns in self.TYPE_MARKERS.items():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return type_name
        
        return 'note'
    
    def _reindex_scenarios(self):
        """Renumber scenario IDs sequentially, removing gaps."""
        # Find all unique scenario IDs that are actually used (not None)
        used_scenarios = sorted(set(u.scenario_id for u in self.units if u.scenario_id))
        
        # Create mapping from old IDs to new sequential IDs
        scenario_map = {}
        for i, old_id in enumerate(used_scenarios, start=1):
            scenario_map[old_id] = f"SCN_{i:03d}"
        
        # Update all units
        for unit in self.units:
            if unit.scenario_id and unit.scenario_id in scenario_map:
                unit.scenario_id = scenario_map[unit.scenario_id]
        
        # Update counter to reflect actual number of used scenarios
        self.scenario_counter = len(used_scenarios)
    
    def generate_units(self):
        """Main orchestration."""
        file_ext = Path(self.source_path).suffix.lower()
        
        if file_ext == '.pdf':
            sections = self.extract_from_pdf()
        elif file_ext == '.docx':
            sections = self.extract_from_docx()
        elif file_ext == '.txt':
            sections = self.extract_from_txt()
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        for page_num, text, section in sections:
            candidates = self.segment_text(text, page_num, section)
            
            for candidate in candidates:
                self.unit_counter += 1
                
                unit = AcceptanceUnit(
                    uas_id=f"UAS_{self.unit_counter:03d}",
                    text=candidate['text'],
                    source_section=candidate.get('section'),
                    source_type=self.classify_source_type(candidate['text']),
                    page=candidate.get('page'),
                    paragraph=candidate.get('paragraph'),
                    source_file=self.source_file,
                    scenario_id=candidate.get('scenario_id'),
                    is_header=candidate.get('is_header', False),
                    actor_role=self._extract_actor_role(candidate['text'], candidate.get('is_header', False))
                )
                
                self.units.append(unit)
        
        # Post-process to ensure sequential scenario IDs with no gaps
        self._reindex_scenarios()
    
 
    def to_json(self, output_path: str):
        """
        Export acceptance units to JSON following the strict schema.
        Saves file inside 'output' folder next to this script.
        """

        import os
        import json
        from datetime import datetime

        # Get directory where this script is located
        base_dir = os.path.dirname(os.path.abspath(__file__))

        # Create output directory inside script folder
        output_dir = os.path.join(base_dir, "output")
        os.makedirs(output_dir, exist_ok=True)

        # Save file inside that directory
        output_path = os.path.join(output_dir, output_path)

        output = {
            "metadata": {
                "total_units": len(self.units),
                "pages_processed": self.pages_processed,
                "source_file": self.source_file,
                "extraction_timestamp": datetime.utcnow().isoformat() + "Z"
            },
            "acceptance_units": [unit.to_dict() for unit in self.units]
        }

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False)

        return output

def ingest_acceptance_stories(input_file: str, output_file: str = "acceptance_units.json") -> Dict:
    """Main entry point for Module 5.1."""
    print(f"\n{'='*70}")
    print("MODULE 5.1: Acceptance Story Ingestion & Structuring")
    print(f"{'='*70}\n")
    
    ingester = AcceptanceStoryIngester(input_file)
    
    print(f"📄 Processing: {ingester.source_file}")
    print(f"🔍 Extracting acceptance units...")
    
    ingester.generate_units()
    
    print(f"\n✅ Extraction Complete")
    print(f"   • Acceptance Units: {len(ingester.units)}")
    print(f"   • Pages Processed: {ingester.pages_processed}")
    print(f"   • Scenarios Detected: {ingester.scenario_counter}")
    print(f"   • Source File: {ingester.source_file}")
    
    type_counts = {}
    for unit in ingester.units:
        type_counts[unit.source_type] = type_counts.get(unit.source_type, 0) + 1
    
    print(f"\n📊 Type Distribution:")
    for type_name, count in sorted(type_counts.items()):
        print(f"   • {type_name}: {count}")
    
    role_counts = {}
    for unit in ingester.units:
        role = unit.actor_role or "none"
        role_counts[role] = role_counts.get(role, 0) + 1
    
    print(f"\n👤 Actor Role Distribution:")
    for role, count in sorted(role_counts.items()):
        print(f"   • {role}: {count}")
    
    header_count = sum(1 for unit in ingester.units if unit.is_header)
    print(f"\n📋 Headers Detected: {header_count}")
    
    result = ingester.to_json(output_file)
    
    print(f"\n💾 Output: {output_file}")
    print(f"{'='*70}\n")
    
    return result


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python ingest_acceptance_stories.py <input_file> [output_file]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else "acceptance_units.json"
    
    ingest_acceptance_stories(input_file, output_file)